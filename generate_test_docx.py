from docx import Document

# 创建一个新的Word文档
doc = Document()

# 添加标题
doc.add_heading('题库测试文件', 0)

# 添加表头行
doc.add_paragraph('知识点,难度,题型,题目内容,选项A,选项B,选项C,选项D,答案')

# 添加题目数据
questions = [
    'Python基础,简单,single,Python的创始人是谁?,Guido van Rossum,Bjarne Stroustrup,James Gosling,Linus Torvalds,A',
    'Python基础,简单,judge,Python是一种解释型语言吗?,True,False,,,# 正确答案是True',
    'Python基础,中等,multiple,以下哪些是Python的关键字?,if,for,class,function,if,class',
    'Python基础,困难,short,请简述Python的装饰器作用,.,.,.,.,装饰器可以在不修改原函数代码的情况下，增加函数的功能'
]

# 将题目数据添加到文档中
for question in questions:
    doc.add_paragraph(question)

# 保存文档
doc.save('test_questions.docx')
print("测试用DOCX文件已生成：test_questions.docx")